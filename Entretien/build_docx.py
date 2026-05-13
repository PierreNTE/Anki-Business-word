"""Génère le document Word complet de préparation entretien Big Tech pour Pierre Logre."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"c:\Users\plogr\Desktop\Entretien\Pierre-LOGRE_Cheat-Sheet-Entretien-BigTech.docx"

# -- Couleurs charte --
NAVY = RGBColor(0x0B, 0x2A, 0x4A)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0xF2, 0xF2, 0xF2)

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Style normal
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_title(text, size=24, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT, after=10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_h1(text):
    add_title(text, size=18, color=NAVY, after=6)
    # Petite barre sous le H1
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("━" * 30)
    run.font.color.rgb = ACCENT
    run.font.size = Pt(8)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT


def add_para(text, bold=False, italic=False, size=10.5, color=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(hdr[i], '0B2A4A')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    # Espace après tableau
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def page_break():
    doc.add_page_break()


# =================================================================
# COUVERTURE
# =================================================================
add_title("CHEAT SHEET ENTRETIEN", size=28, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
add_title("AE / CSM — Cloud & IA · Big Tech", size=16, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Pierre LOGRE")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Customer Success Manager  |  Cloud & AI")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("p.logre@outlook.fr  ·  +33 6 31 33 77 89  ·  linkedin.com/in/pierre-logre  ·  pierrelogre.com")
r.font.size = Pt(10)
r.font.color.rgb = GREY

# Box cibles
doc.add_paragraph().paragraph_format.space_after = Pt(20)
add_para("CIBLES : AWS · Microsoft · Google · Salesforce · Oracle · ServiceNow · NVIDIA · Dell · HPE · Cisco",
         bold=True, size=11, color=NAVY, after=4)
add_para("MANTRA : « Customer Obsession × Data-driven × Bias for Action »",
         italic=True, size=11, color=ACCENT, after=20)

add_h2("📑 Sommaire")
toc = [
    "Partie 1 — Elevator Pitch (FR + EN)",
    "Partie 2 — Vocabulaire Big Tech (FR/EN)",
    "Partie 3 — Mes 3 histoires STAR fondamentales",
    "Partie 4 — Cartographie marché US Cloud & IA (2026)",
    "Partie 5 — Cheat codes par entreprise (AWS / MS / Google / Salesforce)",
    "Partie 6 — Vocabulaire exhaustif (Sales, Tech, C-level)",
    "Partie 7 — Frameworks à citer en entretien",
    "Partie 8 — Argumentaires « killer » & Plan 30-60-90",
    "Partie 9 — Questions à poser en fin d'entretien",
    "Partie 10 — Pièges & Mantras de pré-entretien",
]
for t in toc:
    add_bullet(t)

page_break()

# =================================================================
# PARTIE 1 — ELEVATOR PITCH
# =================================================================
add_h1("PARTIE 1 — Elevator Pitch (2 minutes)")

add_h2("🇫🇷 Version Française")
add_para(
    "Bonjour, je suis Pierre Logre, Account Manager / Customer Success avec plus de 3 ans d'expérience "
    "en vente B2B IT sur les technologies Cloud, SaaS et Infrastructure, acquise chez Sopra Steria et TGS France, "
    "sur les segments Retail et Secteur Public."
)
add_para(
    "Ce qui me différencie, c'est mon ADN hybride : je suis un commercial profondément orienté technique. "
    "Je ne me contente pas de vendre une solution, je la comprends. J'ai d'ailleurs validé des certifications "
    "OVHcloud (IAM, IaaS, PaaS, Enterprise), Numspot, Dynatrace et Anthropic Claude, "
    "et je développe moi-même des outils sur PowerApps et ServiceNow pour gagner en productivité et challenger mes clients."
)
add_para(
    "Concrètement, j'ai généré plus de 650 K€ de revenus à ~18 % de marge, remporté un appel d'offres public majeur — "
    "le CAIH — avec OVHcloud et Numspot face à des acteurs établis, et qualifié 80+ leads sur cycle complet chez TGS "
    "grâce à une prospection multicanale structurée."
)
add_para(
    "Aujourd'hui, je veux mettre cette double casquette Sales + Tech au service d'un acteur comme [AWS/Microsoft/"
    "Google/Salesforce], parce que je crois que le futur de la vente Cloud & IA appartient aux profils capables d'être "
    "de véritables Trusted Advisors — et pas de simples vendeurs de licences."
)

add_h2("🇬🇧 English Version")
add_para(
    "Hi, I'm Pierre Logre, an Account Manager / Customer Success professional with 3+ years of B2B IT sales experience "
    "across Cloud, SaaS and Infrastructure, built at Sopra Steria and TGS France, covering Retail and Public Sector segments."
)
add_para(
    "What sets me apart is my hybrid DNA — I'm a deeply tech-oriented sales professional. I don't just sell a solution, "
    "I understand it. I've earned partner certifications on OVHcloud (IAM, IaaS, PaaS, Enterprise), Numspot, Dynatrace and "
    "Anthropic Claude, and I personally build productivity tools on PowerApps and ServiceNow to challenge my customers and "
    "scale my own workflow."
)
add_para(
    "On the numbers side, I've generated 650K€+ in revenue at ~18% growth margin, won a major public tender — the CAIH — "
    "partnering with OVHcloud and Numspot against established players, and qualified 80+ leads end-to-end at TGS through a "
    "structured multichannel motion."
)
add_para(
    "Today, I want to bring this Sales + Tech combination to a company like [AWS/Microsoft/Google/Salesforce], because I "
    "believe the future of Cloud & AI sales belongs to people who can be true Trusted Advisors — not just license sellers."
)

add_h3("💡 Tips de delivery")
for t in [
    "Ton calme, débit lent — montrer la sérénité d'un closer.",
    "1 chiffre fort toutes les 20 secondes (650K€, 18%, 80+, CAIH).",
    "Terminer sur : « ...and that's exactly why I'm here today. »",
    "Maintenir contact visuel — ne pas lire ses notes.",
]:
    add_bullet(t)

page_break()

# =================================================================
# PARTIE 2 — VOCABULAIRE BIG TECH (10 verbes)
# =================================================================
add_h1("PARTIE 2 — Vocabulaire Big Tech (Top 10)")

vocab10 = [
    ("To spearhead", "Piloter / mener de front", "I spearheaded the GTM strategy for a new IaaS offering, opening 12 enterprise accounts in 6 months."),
    ("Data-driven", "Piloté par la donnée", "My pipeline reviews are 100% data-driven — I track win-rate per stage in Salesforce to forecast within ±5%."),
    ("To navigate ambiguity", "Évoluer dans l'incertitude", "In a public RFP with shifting specs, I navigated ambiguity by aligning OVHcloud and Numspot on a unified PaaS narrative."),
    ("Trusted Advisor", "Conseiller de confiance", "For my SaaS accounts I act as a Trusted Advisor — I challenge the customer's roadmap before pushing any upsell."),
    ("To uncover pain points", "Révéler les irritants", "Through MEDDIC discovery I uncovered pain points on observability that unlocked a Dynatrace SaaS expansion."),
    ("To drive adoption", "Stimuler l'adoption", "I drove adoption of our PaaS sandbox from 15 to 80+ active users in one quarter."),
    ("Cross-functional collaboration", "Collaboration transverse", "I led cross-functional collaboration with Solution Architects and Legal to close a complex IaaS migration."),
    ("To exceed quota / overachieve", "Dépasser son objectif", "I exceeded quota by 130%, generating 650K€+ on Cloud & Infrastructure deals."),
    ("Customer Obsession", "Obsession client", "Customer Obsession means I joined my client's internal steerco — not to sell, but to de-risk their AI rollout."),
    ("To land and expand", "Conquérir puis développer", "Classic land-and-expand: I landed a 30K€ SaaS pilot, then expanded to a 220K€ enterprise license within 9 months."),
]
add_table(["EN — Verbe / Concept", "FR — Équivalent", "Exemple appliqué (IaaS / PaaS / SaaS)"],
          vocab10, col_widths=[4, 4, 9])

add_h3("🔥 Bonus power-words à placer")
add_para(
    "bias for action · ownership · raise the bar · dive deep · earn trust · think big · scrappy · whitespace · "
    "co-sell · workload migration · consumption-based · stickiness · de-risk · future-proof · to articulate · "
    "to displace · to incubate · to productize · to over-index · to dogfood.",
    italic=True
)

page_break()

# =================================================================
# PARTIE 3 — STAR
# =================================================================
add_h1("PARTIE 3 — Mes 3 histoires STAR fondamentales")

add_quote("Règle d'or : toujours dire « I » (pas « we »). Quantifier chaque résultat. Mapper à 1-2 Leadership Principles.")

# STAR 1 — CAIH
add_h2("🏆 STAR #1 — La victoire complexe : Appel d'offres CAIH (OVHcloud × Numspot)")
add_para("Leadership Principles : Deliver Results · Earn Trust · Think Big", bold=True, color=ACCENT)

add_h3("Trame complète (FR)")
add_bullet("Chez Sopra Steria, le CAIH (Centrale d'Achat de l'Informatique Hospitalière) a publié un appel d'offres stratégique "
           "pour un Cloud souverain destiné à 1000+ établissements de santé. Concurrence frontale avec des acteurs hyperscalers installés.",
           bold_prefix="Situation — ")
add_bullet("En tant qu'AE lead sur le dossier, ma mission était de construire la réponse technique et commerciale en orchestrant "
           "deux partenaires majeurs : OVHcloud (IaaS souverain) et Numspot (PaaS SecNumCloud).",
           bold_prefix="Task — ")
add_bullet("J'ai (1) cartographié les critères de notation ligne par ligne, (2) organisé 7 ateliers techniques entre nos trois entités "
           "pour aligner l'architecture cible, (3) rédigé personnellement la note commerciale + le mémoire technique synthétique, "
           "(4) négocié la grille tarifaire pour rester compétitif sans rogner la marge, et (5) défendu l'offre en soutenance orale devant le jury CAIH.",
           bold_prefix="Action — ")
add_bullet("Appel d'offres remporté face à 4 concurrents · ouverture d'un canal de revenus pluriannuel sur la santé publique · reconnaissance "
           "interne comme référent « secteur public Cloud souverain ».",
           bold_prefix="Result — ")

add_h3("🇬🇧 Bullet points (à lire en entretien)")
for line in [
    "SITUATION — Sopra Steria · public RFP from CAIH (French hospital purchasing central) · sovereign cloud for 1000+ hospitals · facing hyperscaler incumbents.",
    "TASK — Lead AE: owned the full bid response · orchestrate two strategic partners (OVHcloud IaaS + Numspot SecNumCloud PaaS).",
    "ACTION — Mapped every scoring criterion line-by-line · ran 7 cross-company technical workshops · personally drafted the commercial offer + exec tech summary · negotiated pricing grid (kept margin) · defended the bid in live oral panel.",
    "RESULT — Won the RFP against 4 competitors · opened a multi-year revenue channel in public healthcare · became internal go-to expert on Sovereign Cloud / Public Sector.",
    "LEADERSHIP PRINCIPLE → Deliver Results + Earn Trust.",
]:
    add_bullet(line)

# STAR 2
add_h2("🎯 STAR #2 — La résilience & la chasse : 80+ leads qualifiés chez TGS France")
add_para("Leadership Principles : Bias for Action · Ownership · Are Right A Lot", bold=True, color=ACCENT)

add_h3("Trame complète (FR)")
add_bullet("Chez TGS France, je rejoins une équipe sur un marché Cloud / SaaS / M365 saturé où le pipeline entrant était quasi-inexistant. "
           "Aucun playbook outbound structuré n'existait sur mon périmètre.",
           bold_prefix="Situation — ")
add_bullet("Objectif personnel : générer mon propre pipeline de A à Z — du sourcing au RDV qualifié — sans dépendre du marketing. "
           "Cible : renouvellement M365 et upsell cybersécurité, ERP, CRM, hardware (TD Synnex) sur ~50 PME.",
           bold_prefix="Task — ")
add_bullet("J'ai (1) construit un ICP précis (taille, secteur, stack, signaux d'achat), (2) mis en place une cadence multicanale "
           "(LinkedIn + email séquencé + cold call + événements + RDV terrain), (3) testé et itéré 3 versions de séquences avec A/B testing "
           "sur l'objet d'email, (4) industrialisé le suivi dans le CRM avec des vues custom de scoring, et (5) tenu une discipline "
           "quotidienne de 30+ touches/jour.",
           bold_prefix="Action — ")
add_bullet("80+ leads qualifiés générés sur cycle complet · taux de transformation >18 % sur les RDV décrochés · méthode reprise "
           "comme template par 2 autres commerciaux de l'équipe.",
           bold_prefix="Result — ")

add_h3("🇬🇧 Bullet points")
for line in [
    "SITUATION — TGS France · saturated Cloud/SaaS/M365 market · near-zero inbound on my territory · no outbound playbook.",
    "TASK — Build my own pipeline end-to-end · sourcing to qualified meeting · no reliance on marketing · ~50 SMB targets for M365 renewals + cyber/ERP/CRM/HW upsell.",
    "ACTION — Defined sharp ICP · multichannel cadence (LinkedIn + sequenced email + cold call + events + field meetings) · A/B tested 3 sequence versions · industrialized CRM tracking with custom scoring views · daily discipline of 30+ touches.",
    "RESULT — 80+ qualified leads, full-cycle · >18% conversion on booked meetings · methodology adopted by 2 teammates.",
    "LEADERSHIP PRINCIPLE → Bias for Action + Ownership.",
]:
    add_bullet(line)

# STAR 3
add_h2("🛠️ STAR #3 — L'innovation Tech : application PowerApps de A à Z chez Metsys")
add_para("Leadership Principles : Invent and Simplify · Customer Obsession · Learn and Be Curious", bold=True, color=ACCENT)

add_h3("Trame complète (FR)")
add_bullet("Chez Metsys, les équipes métier perdaient un temps significatif sur un processus manuel (suivi & planification de capacité) "
           "reposant sur des fichiers Excel partagés — sources d'erreurs, de doublons et de retards.",
           bold_prefix="Situation — ")
add_bullet("Personne n'était officiellement missionné. J'ai pris l'initiative de proposer une solution low-code (PowerApps) plutôt que "
           "d'attendre un projet IT classique long et coûteux. Méthodologie cadrée : Scrum + ADKAR + ITIL.",
           bold_prefix="Task — ")
add_bullet("J'ai (1) mené 5 interviews utilisateurs pour cartographier le vrai irritant, (2) conçu la maquette UX (papier puis Figma), "
           "(3) développé l'application PowerApps de A à Z (formulaires, logique métier, connecteurs Dataverse / SharePoint), "
           "(4) intégré un workflow Power Automate pour notifier les responsables, et (5) formé les utilisateurs (doc + session live).",
           bold_prefix="Action — ")
add_bullet("App déployée en production · temps de traitement significativement réduit · zéro doublon depuis la mise en service · "
           "reconnaissance interne comme profil « Sales-Tech » capable de livrer un outil métier — pas juste de le demander.",
           bold_prefix="Result — ")

add_h3("🇬🇧 Bullet points")
for line in [
    "SITUATION — Metsys · business teams losing hours on a manual Excel-based capacity-planning process · errors, duplicates, delays · no IT project planned.",
    "TASK — Nobody formally assigned · I took the initiative · proposed a low-code (PowerApps) solution · framed with Scrum + ADKAR + ITIL.",
    "ACTION — Ran 5 user interviews to map real pain · designed UX wireframes (paper → Figma) · built the PowerApps app end-to-end (forms, logic, Dataverse/SharePoint connectors) · added a Power Automate workflow for stakeholder notifications · trained users via doc + live session.",
    "RESULT — App shipped to production · processing time significantly reduced · zero duplicates since go-live · earned a reputation as a Sales-Tech hybrid who SHIPS.",
    "LEADERSHIP PRINCIPLE → Invent and Simplify + Customer Obsession.",
]:
    add_bullet(line)

page_break()

# =================================================================
# PARTIE 4 — MARCHÉ
# =================================================================
add_h1("PARTIE 4 — Cartographie du marché US Cloud & IA (2026)")

add_h2("4.1 — Le « Big 3 » des hyperscalers (~65 % du marché Cloud mondial)")
add_table(
    ["Acteur", "Part IaaS/PaaS", "Positionnement", "Force différenciante", "Faiblesse perçue"],
    [
        ["AWS", "~31 %", "The default cloud", "Largeur de catalogue (200+ services), maturité, écosystème re:Invent", "Complexité tarifaire, UX console"],
        ["Microsoft Azure", "~25 %", "The enterprise cloud", "Intégration M365 / Entra ID, co-sell partenaires, exclu OpenAI", "Image legacy sur certaines workloads"],
        ["Google Cloud (GCP)", "~11 %", "The data & AI cloud", "BigQuery, Vertex AI, Gemini, DevEx développeur", "Footprint commercial plus petit, profondeur enterprise plus jeune"],
    ],
    col_widths=[3, 2, 3, 5, 4]
)

add_h2("4.2 — Le « Challenger pack » (SaaS, data, AI, hybride)")
add_table(
    ["Acteur", "Catégorie", "Produit phare", "Angle de vente"],
    [
        ["Salesforce", "CRM / Data Cloud / Agentforce", "Sales Cloud, Data Cloud, Agentforce", "« AI-native CRM » — vente assistée par agents"],
        ["Oracle (OCI)", "IaaS / DB", "Autonomous DB, OCI Gen2", "« Cloud for mission-critical workloads »"],
        ["IBM", "Hybrid Cloud / AI", "watsonx, Red Hat OpenShift", "« Enterprise AI + open hybrid » (régulé)"],
        ["ServiceNow", "SaaS Workflow / ITSM", "Now Platform, Now Assist", "« Platform of platforms » IT/HR/Customer"],
        ["Snowflake", "Data Cloud", "Data Cloud, Cortex AI", "« Data sharing without movement »"],
        ["Databricks", "Data + AI lakehouse", "Mosaic AI, Unity Catalog", "« Lakehouse paradigm » — alternative ouverte à Snowflake"],
        ["Workday / SAP", "SaaS ERP / HCM", "Workday HCM, SAP S/4HANA", "Verticalisation finance & RH"],
    ],
    col_widths=[3, 3.5, 4, 6.5]
)

add_h2("4.3 — Les « Constructeurs » (hardware-meets-cloud)")
add_table(
    ["Acteur", "Rôle clé 2026", "Lien avec les cloudeurs"],
    [
        ["NVIDIA", "Quasi-monopole GPU IA (H100, B200, GB200) — « picks-and-shovels of the AI gold rush »", "Vendu via AWS (P5), Azure (ND), GCP (A3)"],
        ["Dell Technologies", "Serveurs IA, Dell APEX (cloud-as-a-service on-prem)", "Partenaire fort Microsoft & Red Hat"],
        ["HPE", "HPE GreenLake (consumption hybrid), acquisition Juniper", "Concurrent direct des hyperscalers sur l'edge"],
        ["Cisco", "Réseau, sécurité (Splunk), Cisco AI Defense", "Brique infrastructure des hyperscalers"],
        ["Pure Storage / NetApp", "Stockage haute perf pour IA", "Backbone des datalakes"],
        ["Lenovo / Supermicro", "OEM serveurs IA", "Sous-traitants clés de NVIDIA"],
        ["AMD · Intel", "CPU + GPU alternatifs (MI300X, Gaudi 3)", "Tentative de casser le monopole NVIDIA"],
    ],
    col_widths=[3, 7, 7]
)

add_h2("4.4 — Les 8 tendances 2026 à citer en entretien")
trends = [
    ("GenAI → Agentic AI", "On passe des copilotes (assistants) aux agents autonomes : Salesforce Agentforce, Microsoft Copilot Studio, AWS Bedrock Agents, Google Agentspace."),
    ("Sovereign Cloud", "Explosion de la demande UE (RGPD, DORA, NIS2) → opportunités OVHcloud, Numspot, S3NS (Thales × Google), Bleu (Capgemini × Orange × Microsoft)."),
    ("FinOps obligatoire", "Le client veut optimiser sa facture cloud (~40 % de gaspillage moyen) — angle CSM puissant."),
    ("Repatriation partielle", "Certains workloads reviennent on-prem ou hybride (cf. 37signals). L'hybrid cloud revient en force."),
    ("AI Infrastructure crunch", "Pénurie de GPU, files d'attente sur capacité → l'AE qui sécurise des reserved instances devient héros."),
    ("Verticalisation", "AWS for Healthcare, Azure for Industry, Google Public Sector — offres sectorielles santé, finance, secteur public, retail."),
    ("Co-sell & Marketplaces", "80 % des deals enterprise passent par AWS / Azure / Google Marketplace — maîtriser les Private Offers est devenu critique."),
    ("Outcome-based selling", "On ne vend plus du compute, on vend un business outcome mesurable (churn, time-to-market, coût/transaction)."),
]
for h, d in trends:
    add_bullet(d, bold_prefix=h + " — ")

page_break()

# =================================================================
# PARTIE 5 — CHEAT CODES PAR ENTREPRISE
# =================================================================
add_h1("PARTIE 5 — Cheat codes par entreprise cible")

# AWS
add_h2("5.1 — AWS (Amazon Web Services)")
add_table(
    ["Élément", "À mémoriser"],
    [
        ["Leadership Principles (16)", "Customer Obsession · Ownership · Invent and Simplify · Bias for Action · Deliver Results · Earn Trust · Dive Deep · Have Backbone Disagree and Commit · Are Right A Lot · Think Big · Frugality · Hire and Develop the Best · Insist on the Highest Standards · Learn and Be Curious · Strive to be Earth's Best Employer · Success and Scale Bring Broad Responsibility"],
        ["Vocabulaire interne", "Working Backwards · PR/FAQ · Two-pizza team · Day 1 mentality · Mechanism > good intentions · Single-threaded leader"],
        ["Produits stars 2026", "Bedrock (GenAI), Q (assistant), SageMaker, Trainium2 / Inferentia, Outposts (hybride), Marketplace, Connect, Q Developer"],
        ["Format entretien", "Bar Raiser · 100 % comportemental sur les LPs · STAR strict · « Tell me about a time… »"],
        ["Killer question à poser", "« How does this team operationalize Customer Obsession when it conflicts with short-term revenue? »"],
    ],
    col_widths=[4, 13]
)

# Microsoft
add_h2("5.2 — Microsoft")
add_table(
    ["Élément", "À mémoriser"],
    [
        ["Culture / valeurs", "Growth mindset (Satya Nadella) · Customer Obsession · One Microsoft · Diversity & Inclusion · Make a Difference"],
        ["Vocabulaire interne", "Co-sell · MACC (Microsoft Azure Consumption Commitment) · Solution Area · ACR (Azure Consumed Revenue) · FY26 H2 (fiscal year)"],
        ["Produits stars 2026", "Azure OpenAI, Copilot (M365, Studio, Security), Fabric (data), Dynamics 365, Entra, Defender, GitHub Copilot Enterprise"],
        ["Format entretien", "Mix comportemental + cas client · « How would you grow this account? » · forte attente sur partner ecosystem"],
        ["Killer question à poser", "« How is the Copilot attach rate evolving on the Azure base, and where does an AE add the most leverage? »"],
    ],
    col_widths=[4, 13]
)

# Google
add_h2("5.3 — Google Cloud (GCP)")
add_table(
    ["Élément", "À mémoriser"],
    [
        ["Culture / valeurs", "Googleyness · GCA (General Cognitive Ability) · Leadership · Role-related knowledge — 4 piliers d'évaluation"],
        ["Vocabulaire interne", "10x thinking · OKRs · Smart creatives · Launch and iterate · TPM (Technical Program Manager)"],
        ["Produits stars 2026", "Vertex AI, Gemini (1.5 / 2.0), BigQuery, Anthos, Apigee, Google Distributed Cloud (sovereign)"],
        ["Format entretien", "Études de cas analytiques · forte attente sur data fluency · structurer en frameworks"],
        ["Killer question à poser", "« How do you see Gemini's enterprise traction vs. Azure OpenAI in regulated EU industries? »"],
    ],
    col_widths=[4, 13]
)

# Salesforce
add_h2("5.4 — Salesforce")
add_table(
    ["Élément", "À mémoriser"],
    [
        ["Culture / valeurs (Ohana)", "Trust · Customer Success · Innovation · Equality · Sustainability"],
        ["Vocabulaire interne", "Trailblazer · V2MOM (Vision, Values, Methods, Obstacles, Measures) · Customer 360 · Dreamforce · Ohana"],
        ["Produits stars 2026", "Agentforce (agents IA), Data Cloud, Sales Cloud, Service Cloud, Slack, Tableau, MuleSoft, Einstein"],
        ["Format entretien", "Storytelling client fort · démontrer values alignment · cas de land-and-expand"],
        ["Killer question à poser", "« How is Agentforce changing the conversation with CIOs vs. traditional Sales Cloud renewals? »"],
    ],
    col_widths=[4, 13]
)

page_break()

# =================================================================
# PARTIE 6 — VOCABULAIRE EXHAUSTIF
# =================================================================
add_h1("PARTIE 6 — Vocabulaire exhaustif")

add_h2("6.1 — Verbes d'action premium (Sales / GTM)")
add_table(
    ["EN", "FR", "Quand l'utiliser"],
    [
        ["To spearhead", "Piloter / mener", "Initiative dont on est leader"],
        ["To orchestrate", "Orchestrer", "Multi-stakeholders / partenaires"],
        ["To champion", "Porter en interne", "Défendre une cause / un produit"],
        ["To unlock", "Débloquer", "Lever un frein business"],
        ["To uncover", "Mettre au jour", "Discovery / pain points"],
        ["To accelerate", "Accélérer", "Time-to-value / cycle de vente"],
        ["To de-risk", "Sécuriser", "Réduire le risque d'un projet"],
        ["To future-proof", "Pérenniser", "Architecture évolutive"],
        ["To benchmark", "Comparer", "Vs. concurrents ou peers"],
        ["To articulate", "Formuler clairement", "Valeur, vision, ROI"],
        ["To quantify", "Chiffrer", "Business case"],
        ["To navigate ambiguity", "Évoluer dans le flou", "Contexte incertain"],
        ["To rally", "Mobiliser", "Une équipe ou un sponsor"],
        ["To position", "Positionner", "Une offre face au marché"],
        ["To displace", "Évincer (concurrent)", "Contexte compétitif"],
        ["To incubate", "Faire émerger", "Nouveau use case"],
        ["To productize", "Industrialiser", "D'un POC à un produit"],
        ["To monetize", "Monétiser", "Pricing strategy"],
        ["To over-index", "Sur-investir sur", "« We over-index on customer outcomes »"],
        ["To dogfood", "Utiliser son propre produit", "Crédibilité tech"],
    ],
    col_widths=[4, 5, 8]
)

add_h2("6.2 — Concepts business / commerciaux")
add_table(
    ["Terme EN", "Définition courte", "Usage"],
    [
        ["ACV (Annual Contract Value)", "Valeur annuelle du contrat", "KPI AE #1"],
        ["ARR (Annual Recurring Revenue)", "Revenu récurrent annuel", "KPI SaaS"],
        ["NRR / GRR", "Net / Gross Retention Rate", "KPI CSM #1"],
        ["TCV (Total Contract Value)", "Valeur totale multi-années", "Deals stratégiques"],
        ["CAC / LTV", "Coût d'acquisition / valeur vie", "Unit economics"],
        ["Pipeline coverage", "Couverture pipeline (× quota)", "3-4× = healthy"],
        ["Win rate", "Taux de victoire", "Forecast accuracy"],
        ["Time-to-close", "Durée de cycle", "Vélocité"],
        ["Sandbagging", "Sous-estimer son forecast", "À ne PAS faire"],
        ["Quota attainment", "% d'atteinte d'objectif", "KPI personnel"],
        ["Land and expand", "Conquérir puis développer", "Stratégie de compte"],
        ["Whitespace", "Comptes / produits non couverts", "Cible upsell"],
        ["Cross-sell / Upsell", "Vente croisée / montée en gamme", "CSM motion"],
        ["Co-sell", "Vente conjointe avec partenaire", "Microsoft / AWS playbook"],
        ["Marketplace private offer", "Offre négociée sur marketplace cloud", "Levier de closing"],
        ["MEDDIC / MEDDPICC", "Framework de qualification", "À citer en entretien"],
        ["BANT", "Budget Authority Need Timing", "Framework legacy mais connu"],
        ["Command of the Message", "Force Management — discours valeur", "Méthode reconnue"],
        ["Champion / Economic Buyer", "Allié / décideur financier", "Stakeholder map"],
        ["Mobilizer", "Profil moteur du changement", "CEB / Gartner"],
        ["Power user", "Utilisateur expert produit", "CSM leverage"],
        ["Sponsor / Detractor", "Soutien / opposant", "Cartographie politique"],
        ["POC / Pilot", "Preuve de concept / pilote", "Étape de cycle"],
        ["Bake-off", "Compétition technique côte à côte", "Vendor selection"],
        ["Greenfield", "Compte neuf", "Hunting"],
        ["Brownfield", "Compte existant", "Farming"],
        ["Logo win", "Conquête d'un nouveau client", "KPI hunter"],
        ["Stickiness", "Adhérence du produit", "CSM angle"],
        ["Net new logo", "Nouveau client net", "Acquisition"],
        ["QBR (Quarterly Business Review)", "Revue trimestrielle client", "CSM cadence"],
        ["Health score", "Score de santé d'un compte", "CSM dashboard"],
        ["Time to value (TTV)", "Délai de génération de valeur", "Onboarding KPI"],
    ],
    col_widths=[5, 6, 6]
)

add_h2("6.3 — Concepts techniques Cloud / IA à maîtriser")
add_table(
    ["Terme", "Définition courte"],
    [
        ["IaaS / PaaS / SaaS / FaaS", "Couches d'abstraction cloud (Infrastructure, Plateforme, Software, Function)"],
        ["Multi-cloud / Hybrid cloud", "Plusieurs clouds / cloud + on-prem"],
        ["Lift-and-shift", "Migration sans refactoring"],
        ["Re-platform / Re-factor", "Migration avec adaptation / refonte"],
        ["Workload migration", "Bascule d'une charge applicative"],
        ["Containerization", "Conteneurs (Docker, Kubernetes)"],
        ["Serverless", "Sans gestion serveur (Lambda, Functions)"],
        ["Edge computing", "Calcul en périphérie"],
        ["Data lakehouse", "Datalake + datawarehouse unifiés"],
        ["Vector database", "DB pour embeddings (Pinecone, pgvector)"],
        ["RAG (Retrieval-Augmented Generation)", "LLM + base de connaissances"],
        ["Fine-tuning", "Spécialisation d'un modèle"],
        ["Prompt engineering", "Art du prompt"],
        ["Inference / Training", "Exécution / entraînement d'un modèle"],
        ["Tokens", "Unité de facturation LLM"],
        ["Context window", "Mémoire d'un LLM"],
        ["Foundation model", "Modèle généraliste (GPT, Claude, Gemini)"],
        ["Agentic AI", "IA capable d'actions autonomes"],
        ["Guardrails", "Garde-fous éthiques / sécurité"],
        ["MLOps / LLMOps", "DevOps appliqué à l'IA"],
        ["Observability", "Logs + métriques + traces (Dynatrace, Datadog)"],
        ["SRE (Site Reliability Engineering)", "Fiabilité système"],
        ["SLA / SLO / SLI", "Niveaux de service"],
        ["Zero Trust", "Architecture sécurité moderne"],
        ["SASE / SSE", "Sécurité réseau cloud"],
        ["FinOps", "Optimisation des coûts cloud"],
        ["Consumption-based pricing", "Facturation à l'usage"],
        ["Reserved instances / Savings plans", "Engagements pour réduire la facture"],
        ["Sovereign cloud", "Cloud souverain (juridiction UE)"],
        ["Air-gapped", "Isolé d'internet (défense, santé)"],
        ["IAM", "Identity & Access Management"],
        ["DORA / NIS2 / RGPD", "Réglementations UE clés"],
    ],
    col_widths=[6, 11]
)

add_h2("6.4 — Expressions « C-level » pour parler comme un Trusted Advisor")
expressions = [
    "« Let me play this back to make sure I understood… » (reformulation)",
    "« What does success look like 12 months from now? » (vision)",
    "« Help me understand the cost of inaction here. » (urgence)",
    "« Who else needs to be in the room to make this decision? » (champion + EB)",
    "« If budget weren't a constraint, what would you change first? » (révéler la vraie priorité)",
    "« I want to be the easiest vendor you've ever worked with. » (positioning)",
    "« Let's pressure-test this assumption together. » (data-driven)",
    "« I'd rather lose the deal than misrepresent the fit. » (earn trust)",
    "« How are you measured on this initiative? » (économique du sponsor)",
    "« What's the ripple effect across your org if we get this right? » (think big)",
    "« Where are you on the maturity curve vs. your peers? » (benchmarking)",
    "« I see three paths forward — let me walk you through trade-offs. » (advisor mode)",
]
for e in expressions:
    add_bullet(e)

add_h2("6.5 — Anti-patterns à BANNIR")
add_table(
    ["❌ À éviter", "✅ Remplacer par"],
    [
        ["« I think / maybe / kind of »", "« Based on the data, I'm confident that… »"],
        ["« We just… »", "« We deliberately… »"],
        ["« It's not my fault »", "« I owned the outcome, here's what I learned… »"],
        ["« The customer doesn't get it »", "« I haven't found the right way to articulate value yet »"],
        ["« Closed a deal » (passif)", "« Drove the deal end-to-end »"],
        ["« Helped the team »", "« Led / spearheaded / owned… »"],
        ["« Good results »", "« +130 % quota attainment / 650K€ ACV »"],
        ["« We did X »", "« I did X » (ownership)"],
    ],
    col_widths=[7, 10]
)

page_break()

# =================================================================
# PARTIE 7 — FRAMEWORKS
# =================================================================
add_h1("PARTIE 7 — Frameworks à citer en entretien")

add_table(
    ["Framework", "Acronyme / Détails", "Quand le citer"],
    [
        ["MEDDPICC", "Metrics · Economic Buyer · Decision Criteria · Decision Process · Paper Process · Identified Pain · Champion · Competition", "Question sur qualification"],
        ["Challenger Sale", "Teach · Tailor · Take Control", "Profil de vendeur moderne"],
        ["SPIN Selling", "Situation · Problem · Implication · Need-payoff", "Discovery"],
        ["Sandler 7-Up", "Bonding · Up-front contracts · Pain · Budget · Decision · Fulfillment · Post-sell", "Cycle de vente complet"],
        ["Value Selling Framework", "Force Management — Command of the Message", "Pitch produit"],
        ["GAP Selling", "Current State → Future State → Cost of inaction", "Discovery moderne"],
        ["Customer-Centric Selling", "Aligner sur le buyer's journey", "Generic-modern"],
        ["Working Backwards (Amazon)", "Press release first → product", "Question vision / innovation"],
        ["DACI / RACI", "Driver · Approver · Contributor · Informed", "Question gouvernance"],
        ["OKR", "Objectives & Key Results", "Question pilotage"],
        ["RICE", "Reach · Impact · Confidence · Effort", "Prioriser un portefeuille"],
        ["AARRR (Pirate Metrics)", "Acquisition · Activation · Retention · Referral · Revenue", "Funnel SaaS"],
    ],
    col_widths=[4, 8, 5]
)

page_break()

# =================================================================
# PARTIE 8 — ARGUMENTAIRES KILLER + 30-60-90
# =================================================================
add_h1("PARTIE 8 — Argumentaires « killer » & Plan 30-60-90")

add_h2("8.1 — 5 angles « killer » pour parler aux Big Tech")

add_h3("Pourquoi un profil « Sales × Tech » comme moi est rare")
add_quote("« Most reps can quote a pricing tier. Few can actually demo a workload. I do both — and that's how I shorten cycles by 20-30 %. »")

add_h3("Pourquoi je veux quitter les SSII / Reseller pour un éditeur / hyperscaler")
add_quote("« In a reseller I was an integrator's voice. At [AWS/Microsoft/Google], I'd be the source — closer to the product, "
          "closer to the roadmap, and able to shape outcomes for customers at a scale I can't reach today. »")

add_h3("Pourquoi le moment est unique (timing)")
add_quote("« We're at an inflection point — AI is rewriting the enterprise stack. The companies that win the next 5 years are "
          "the ones with reps who can hold both a CFO conversation and a CTO conversation in the same meeting. That's exactly "
          "the seat I want. »")

add_h3("Pourquoi je serai ramp-up rapide")
add_quote("« I've already taken the certifications path on my own time — OVHcloud, Numspot, Dynatrace, Anthropic Claude. "
          "I don't wait for enablement; I build my own. Give me 90 days and I'll be billable on quota. »")

add_h3("Pourquoi je serai différent une fois dans le rôle")
add_quote("« I build my own productivity tools — PowerApps, ServiceNow flows. While reps are stuck in CRM hygiene, I'll be "
          "automating mine and spending 30 % more time in front of customers. »")

add_h2("8.2 — Plan 30-60-90 jours (à dégainer en fin d'entretien)")
add_table(
    ["Phase", "Objectifs concrets"],
    [
        ["0-30 jours — Learn",
         "Onboarding produit · certifs internes · shadow 5 deals · cartographier mon territoire · 1:1 avec Solution Engineer, SDR, Partner Manager, CSM, Marketing"],
        ["30-60 jours — Engage",
         "Premier pipeline build · 20+ comptes ciblés · 3 events / executive briefings · alignement avec partner ecosystem · premier MEDDPICC complet"],
        ["60-90 jours — Deliver",
         "Premiers deals en stage 3+ · plan de compte sur top 5 · POC lancé · forecast crédible présenté au manager · premier deal closed-won"],
    ],
    col_widths=[5, 12]
)

page_break()

# =================================================================
# PARTIE 9 — QUESTIONS À POSER
# =================================================================
add_h1("PARTIE 9 — Questions intelligentes à poser en fin d'entretien")

add_h2("Sur le rôle")
for q in [
    "« What separates a top performer from an average AE on this team in their first year? »",
    "« What does the pipeline coverage ratio look like today on this territory? »",
    "« How is the patch defined — vertical, geo, named accounts? »",
]:
    add_bullet(q)

add_h2("Sur l'équipe")
for q in [
    "« How does this team collaborate with Solution Engineers and Partner Managers — friction points? »",
    "« What's the tenure of the top performers and what kept them here? »",
]:
    add_bullet(q)

add_h2("Sur la stratégie")
for q in [
    "« Where do you see the biggest whitespace in EMEA for this product line in the next 18 months? »",
    "« How is the company responding to the Agentic AI shift — internally and in GTM? »",
    "« What's the one thing you wish you'd known before joining this team? » (humain, désarmant)",
]:
    add_bullet(q)

add_h2("Sur le développement")
for q in [
    "« What's the typical path from AE to senior AE / strategic AE here? »",
    "« How does enablement work post-onboarding — formal or self-driven? »",
]:
    add_bullet(q)

page_break()

# =================================================================
# PARTIE 10 — PIÈGES & MANTRAS
# =================================================================
add_h1("PARTIE 10 — Pièges classiques & Mantras de pré-entretien")

add_h2("10.1 — 7 pièges & antidotes")
add_table(
    ["Piège", "Antidote"],
    [
        ["« Tell me about a failure » — réponse trop lisse", "Choisir un VRAI échec · décrire ce qu'on a appris · montrer le changement de comportement"],
        ["Syndrome du « We »", "S'entraîner à dire « I » à voix haute avant l'entretien"],
        ["Pas de chiffres", "Avoir 5 KPIs en tête (€, %, durée, volume, ratio)"],
        ["Pas de questions à la fin", "En préparer 10, en poser 3"],
        ["Mauvaise lecture du LP (AWS)", "Mapper chaque histoire à 2 LPs MAX"],
        ["Critiquer un ancien employeur", "Reformuler en « What I learned that I want to apply differently… »"],
        ["Sur-vendre une compétence non maîtrisée", "Préférer « I'm building this skill — here's my plan »"],
    ],
    col_widths=[7, 10]
)

add_h2("10.2 — Mantras à se répéter avant chaque entretien")
mantras = [
    "« They're not interviewing me. We're testing fit — both ways. »",
    "« Every answer = one story + one number + one lesson. »",
    "« Calm voice, slow pace, sharp data. »",
    "« Customer Obsession is not a slogan — it's a verb. »",
    "« I'm not selling myself. I'm previewing how I'll sell their product. »",
]
for m in mantras:
    add_quote(m)

add_h2("10.3 — Checklist H-1 avant l'entretien")
checklist = [
    "Recherche entreprise : 3 actualités récentes · 1 produit phare · 1 concurrent",
    "Glassdoor : questions fréquentes du rôle visé",
    "STAR ready : capable de raconter chaque histoire en 90 sec max",
    "Questions à poser : 3 questions de fin (1 rôle · 1 équipe · 1 stratégie)",
    "Chiffres-clés en tête : 650K€ · 18 % marge · 80+ leads · CAIH win · 4 certifs",
    "Closing : « Based on what you shared, I'm even more convinced this is the right fit. What are the next steps? »",
    "Tenue : business sobre · webcam à hauteur des yeux · fond neutre",
    "Tech check : micro, caméra, lien de visio testés 15 min avant",
]
for c in checklist:
    add_bullet(c)

# =================================================================
# FOOTER
# =================================================================
doc.add_paragraph().paragraph_format.space_before = Pt(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Document de préparation personnel · Pierre LOGRE · 2026 —")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

doc.save(OUT)
print(f"OK — généré : {OUT}")
